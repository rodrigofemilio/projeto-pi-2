from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from PIL import Image, ImageTk
# Pip install pmw
import Pmw
from datetime import date, timedelta
# Pip install sqlite3
import sqlite3
import webbrowser
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont 
from reportlab.platypus import SimpleDocTemplate, Image
from docx import Document
from docx2pdf import convert
# Pip install tkinterweb
import tkinterweb as tkweb
import sys

root = Tk()
Pmw.initialise(root)

class bd():
    def variaveis_f1(self):
        # Pegar dados das Entry's
        self.codigo = self.f1_entry_cod.get()
        self.nome_aluno = self.f1_entry_aluno.get()
        self.serie_aluno = self.f1_variavel_serie.get()
        self.nome_livro = self.f1_entry_livro.get()
        self.data_retira = self.format_data_retirada
        self.data_entrega = self.format_data_entrega
    def variaveis_imp2(self):
        self.e_c.get()
        self.e_n.get()
        self.e_s.get()
        self.e_s.get()
        self.e_l.get()
        self.e_d1.get()
        self.e_d2.get()
    def conecta_bd_f1(self):
        self.conn = sqlite3.connect('ARQUIVOS\BD\livros.bd'); print('Conectando ao Banco de Dados')
        self.cursor = self.conn.cursor()
    def desconecta_bd_f1(self):
        self.conn.close(); print('Desconectando o Banco de Dados')
    def montatabelas_f1(self):
        self.conecta_bd_f1()
        # Criando a tabela da página de inicio
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS alunos (
                cod             INTEGER   PRIMARY KEY,
                nomes_alunos    CHAR (40) REFERENCES identidade (nome),
                series_alunos   CHAR (10) REFERENCES identidade (serie),
                nomes_livros    CHAR (20) REFERENCES livros (nome_l),
                data_retirada   CHAR (15),
                data_entregada  CHAR (15),
            FOREIGN KEY (
                nomes_alunos,
                series_alunos
                        )
            REFERENCES identidade (nome,serie),
                            
            FOREIGN KEY (
                nomes_livros
                        )
            REFERENCES livros (nome_l)  
            );
        ''')
        self.conn.commit(); print('Banco de Dados 1 criado (Coleta)')
        # Criando a tabela do nome do aluno
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS identidade (
                cod_1           INTEGER   PRIMARY KEY,
                nome            CHAR (40) REFERENCES alunos (nomes_alunos),
                serie           CHAR (10) REFERENCES alunos (series_alunos),
                telefone        CHAR (20),
                identificacao   CHAR (20),
                email           CHAR (40),
                mae             CHAR (40),
                pai             CHAR (40),
            FOREIGN KEY (
                nome
                        )
            REFERENCES alunos (nomes_alunos),
            FOREIGN KEY (
                serie
                        )
            REFERENCES alunos (series_alunos) 
            );
        ''')
        self.conn.commit(); print('Banco de Dados 2 criado (Alunos)')
        # Criando a tabela de livros
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS livros (
                cod_2           INTEGER   PRIMARY KEY,
                nome_l          CHAR (40) REFERENCES alunos (nomes_livros),
                ano             CHAR (10),
                autor           CHAR (20),
                genero          CHAR (20),
                numero          CHAR (6),
            FOREIGN KEY (
                nome_l
                        )
            REFERENCES alunos (nomes_livros) 
            );
        ''')
        self.conn.commit(); print('Banco de Dados 3 criado (Livros)')
        # Criando a tabela de escolas
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS escolas(
                cod_3 INTEGER PRIMARY KEY, 
                nome_e CHAR(30),
                rua CHAR(20),
                numero CHAR(6),
                bairro CHAR(10),
                telefone_e CHAR(15),
                email_e CHAR(30)
            );
        ''')
        self.conn.commit(); print('Banco de Dados 4 criado (Escolas)')
        self.desconecta_bd_f1()

class pdf(bd):
    def abrir_comprovante(self):
        webbrowser.open(f'ARQUIVOS\PDF\COMPROVANTE\comprovante_{(self.nome_aluno)}.pdf')
    def geracomprovante(self):
        # Chamar as Variaveis
        self.variaveis_f1()
        self.c1 = canvas.Canvas(f'ARQUIVOS\PDF\COMPROVANTE\comprovante_{(self.nome_aluno)}.pdf')
        ############################ FOLHA 1 PARTE DO ALUNO
        # Titulo 1 na folha
        self.c1.setFont('Helvetica-Bold', 24)
        self.c1.drawString(80, 650, 'Comprovante')
        # Nome do aluno na parte do aluno 
        self.c1.setFont('Helvetica', 14)
        self.c1.drawString(30, 600, f'Nome do aluno:  {(self.nome_aluno)}')
        self.c1.drawString(25, 595, '______________________________________________________________________')
        # Serie do aluno na parte do professor
        self.c1.setFont('Helvetica', 14)
        self.c1.drawString(420, 570, f'Serie do aluno:  {(self.serie_aluno)}')
        self.c1.drawString(415, 566, '____________________')
        # Dia em que o livro foi pego (Aluno)
        self.c1.setFont('Helvetica', 14)
        self.c1.drawString(30, 540, f'Data de inicio:  {(self.data_retira)}')
        self.c1.drawString(25, 536, '__________________________')
        # Dia em que o livro precisa ser devolvido (Aluno)
        self.c1.setFont('Helvetica', 14)
        self.c1.drawString(30, 480, f'Data de entrega:  {(self.data_entrega)}')
        self.c1.drawString(25, 476, '__________________________')
        # Identificação da escola
        self.c1.setFont('Helvetica-Bold', 18)
        self.c1.drawString(50,780, 'E.M.E.F. Professora Yolanda Jorge')
        self.c1.setFont('Helvetica', 12)
        self.c1.drawString(35, 763, 'Rua Mexico, n° - 868, Jardim das Americas, São Simão - SP')
        self.c1.setFont('Helvetica', 16)
        self.c1.drawString(40,745, 'Telefone: (16) 3984-9124')
        # Logo da Univesp no PDF do comprovante
        univesp_logo = PhotoImage(file="ARQUIVOS\IMAGENS\completo_colorido.png")
        self.c1.drawImage("ARQUIVOS\IMAGENS\completo_colorido.png", 380, 735, width=170, height=60)
        self.c1.drawImage("ARQUIVOS\IMAGENS\completo_colorido.png", 30, 55, width=170, height=60)
        ############################ FOLHA 1 PARTE DO PROFESSOR
        # Linha de separação
        self.c1.setFont('Helvetica', 14)
        self.c1.drawString(16, 380, '_ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _')
        self.c1.drawString(16, 381, ' _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _')
        # Titulo Folha do Professor
        self.c1.setFont('Helvetica-Bold', 24)
        self.c1.drawString(80, 340, 'Comprovante de Emprestimo')
        # Nome do aluno na parte do professor
        self.c1.setFont('Helvetica', 14)
        self.c1.drawString(30, 300, f'Nome do aluno:  {(self.nome_aluno)}')
        self.c1.drawString(25, 295, '______________________________________________________________________')
        # Serie do aluno na parte do professor
        self.c1.setFont('Helvetica', 14)
        self.c1.drawString(420, 270, f'Serie do aluno:  {(self.serie_aluno)}')
        self.c1.drawString(415, 266, '____________________')
        # Dia em que o livro foi pego (Professor)
        self.c1.setFont('Helvetica', 14)
        self.c1.drawString(30, 240, f'Data de inicio:   {(self.data_retira)}')
        self.c1.drawString(25, 236, '__________________________')
        # Dia em que o livro precisa ser devolvido (Professor)
        self.c1.setFont('Helvetica', 14)
        self.c1.drawString(30, 180, f'Data de entrega:   {(self.data_entrega)}')
        self.c1.drawString(25, 176, '__________________________')
        # Identificação do aluno e da professora que fez o emprestimo



        # Salvar folha e abrir no navegador
        self.c1.showPage()
        self.c1.save()
        self.abrir_comprovante()
    def abrir_relatorio(self):
        webbrowser.open(f'ARQUIVOS\PDF\RELATORIO\\relatorio_{(self.e_c.get())}.pdf')
    def gerarelatorio(self):
        # Variaveis
        self.variaveis_imp2()
        # Criar PDF
        self.c2 = canvas.Canvas(f'ARQUIVOS\PDF\RELATORIO\\relatorio_{(self.e_c.get())}.pdf')
        # Titulo PDF
        self.c2.setFont('Helvetica-Bold', 24)
        self.c2.drawString(80, 750, 'Relatório do Livro')
        # Código do aluno
        self.c2.setFont('Helvetica', 14)
        self.c2.drawString(30, 700, f'Código do aluno:  {(self.e_c.get())}')
        self.c2.drawString(25, 695, '______________________________________________________________________')
        # Nome do aluno
        self.c2.setFont('Helvetica', 14)
        self.c2.drawString(30, 670, f'Nome do aluno:  {(self.e_n.get())}')
        self.c2.drawString(25, 665, '______________________________________________________________________')
        # Serie do aluno
        self.c2.setFont('Helvetica', 14)
        self.c2.drawString(30, 640, f'Ano do aluno:  {(self.e_s.get())}')
        self.c2.drawString(25, 635, '______________________________________________________________________')
        # Nome do livro
        self.c2.setFont('Helvetica', 14)
        self.c2.drawString(30, 610, f'Nome do livro:  {(self.e_l.get())}')
        self.c2.drawString(25, 605, '______________________________________________________________________')
        # Datas
        self.c2.setFont('Helvetica', 14)
        self.c2.drawString(30, 580, f'Data de retirada do livro:  {(self.e_d1.get())}')
        self.c2.drawString(25, 575, '______________________________________________________________________')
        self.c2.drawString(30, 550, f'Data de entrega do livro:  {(self.e_d2.get())}')
        self.c2.drawString(25, 545, '______________________________________________________________________')
        # Anotações
        self.c2.setFont('Helvetica-Bold', 24)
        self.c2.drawString(80, 495, 'Anotações')
        # Linhas para anotações
        self.c2.setFont('Helvetica', 14)
        self.c2.drawString(25, 445, '______________________________________________________________________')
        self.c2.drawString(25, 425, '______________________________________________________________________')
        self.c2.drawString(25, 405, '______________________________________________________________________')
        self.c2.drawString(25, 385, '______________________________________________________________________')
        self.c2.drawString(25, 365, '______________________________________________________________________')
        self.c2.drawString(25, 345, '______________________________________________________________________')
        self.c2.drawString(25, 325, '______________________________________________________________________')
        self.c2.drawString(25, 305, '______________________________________________________________________')
        self.c2.drawString(25, 285, '______________________________________________________________________')
        self.c2.drawString(25, 265, '______________________________________________________________________')
        self.c2.drawString(25, 245, '______________________________________________________________________')
        self.c2.drawString(25, 225, '______________________________________________________________________')
        self.c2.drawString(25, 205, '______________________________________________________________________')
        self.c2.drawString(25, 185, '______________________________________________________________________')
        self.c2.drawString(25, 165, '______________________________________________________________________')
        self.c2.drawString(25, 145, '______________________________________________________________________')
        self.c2.drawString(25, 125, '______________________________________________________________________')
        self.c2.drawString(25, 105, '______________________________________________________________________')
        # Logo Univesp
        self.c2.drawImage('ARQUIVOS\IMAGENS\completo_colorido.png', 370, 720, width=170, height=60)
        # Salvar folha e abrir no navegador
        self.c2.showPage()
        self.c2.save()
        self.abrir_relatorio()

class Funcs(pdf,bd):
    def limpar_f1(self):
        self.f1_entry_aluno.delete(0, END)
        self.f1_entry_livro.delete(0, END)
        self.f1_entry_cod.delete(0, END)
        self.entry_data_entrega.delete(0, END)
        self.entry_data_retirada.delete(0, END)
        self.f1_entry_serie.delete(0, END)
    def limpa_f2(self):
        self.f2_entry_aluno.delete(0, END)
        self.f2_entry_serie.delete(0, END)
        self.f2_entry_livro.delete(0, END)
    def cadastrar_f1(self):
        # Chamar variaveis
        self.variaveis_f1()
        if self.f1_entry_aluno.get()=='' or self.f1_entry_livro.get()=='' or self.f1_variavel_serie.get()=='Selecionar':
            messagebox.showinfo(title='ERRO!', message='Insira todos os dados!')
            return
        # Conectar ao BD
        self.conecta_bd_f1()
        # Função de colocar no BD
        self.cursor.execute(''' 
                    INSERT INTO alunos 
                        (nomes_alunos, 
                        series_alunos, 
                        nomes_livros, 
                        data_retirada, 
                        data_entregada)
                    VALUES (?,?,?,?,?)''', 
                        (self.nome_aluno, 
                        self.serie_aluno, 
                        self.nome_livro, 
                        self.data_retira, 
                        self.data_entrega))
        # Funções
        self.conn.commit()
        self.desconecta_bd_f1()
        self.seleciona_f1()
        self.seleciona_f2()
        self.limpar_f1()
    def seleciona_f1(self):
        #Selecionar a Treeview
        self.inser_bd.delete(*self.inser_bd.get_children())
        self.atrasa_bd.delete(*self.atrasa_bd.get_children())
        # Conectar o BD
        self.conecta_bd_f1()
        # Criar variavel sem o self
        lista1 = self.cursor.execute('''  
                            SELECT cod, 
                                nomes_alunos, 
                                series_alunos, 
                                nomes_livros, 
                                data_retirada, 
                                data_entregada 
                            FROM alunos ORDER BY data_entregada ASC; 
                                    ''')
        # Adicionar na Treeview
        for i in lista1:
            self.inser_bd.insert("", END, values=i)
            self.atrasa_bd.insert("", END, values=i)
        # Desconectar o BD
        self.desconecta_bd_f1()
    def seleciona_f2(self):
        #Selecionar a Treeview
        self.pesq_bd.delete(*self.pesq_bd.get_children())
        # Conectar o BD
        self.conecta_bd_f1()
        # Criar variavel sem o self
        lista2 = self.cursor.execute('''  
                            SELECT cod, 
                                nomes_alunos, 
                                series_alunos, 
                                nomes_livros, 
                                data_retirada, 
                                data_entregada 
                            FROM alunos ORDER BY nomes_alunos ASC; 
                                    ''')
        # Adicionar na Treeview
        for i in lista2:
            self.pesq_bd.insert("", END, values=i)
        # Desconectar o BD
        self.desconecta_bd_f1()
    def excluir_f1(self):
        # Chamar Variaveis
        self.variaveis_f1()
        # Conectar o BD
        self.conecta_bd_f1()
        # Função de deletar do BD
        try:
            self.cursor.execute('''DELETE FROM alunos WHERE cod = {} '''.format(self.codigo))
            self.conn.commit()
        except:
            messagebox.showinfo(title='ERRO!',message='Selecione com 2 cliques para ser apagado!')
        # Desconectar o BD
        self.desconecta_bd_f1()
        # Limpar as Entrys
        self.limpar_f1()
        # Atualizar a lista
        self.seleciona_f1()
        self.seleciona_f2()
    def alterar_f1(self):
        # Variaveis
        self.variaveis_f1()
        # Conectar ao BD
        self.conecta_bd_f1()
        # Código de alterar os dados
        self.cursor.execute('''
                    UPDATE alunos 
                        SET nomes_alunos = ?, 
                        series_alunos = ?, 
                        nomes_livros = ?, 
                        data_retirada = ?, 
                        data_entregada = ? 
                    WHERE cod = ? 
                    ''', 
                        (self.nome_aluno, 
                        self.serie_aluno, 
                        self.nome_livro, 
                        self.data_retira, 
                        self.data_entrega, 
                        self.codigo))
        self.conn.commit()
        # Desconectar o BD
        self.desconecta_bd_f1()
        # Atualizar a Tabela
        self.seleciona_f1()
        self.seleciona_f2()
        # Limpar as Entrys
        self.limpar_f1()
    def pesquisar_f2(self):
        # Chamar o BD
        self.conecta_bd_f1()
        # Limpar a lista
        self.pesq_bd.delete(*self.pesq_bd.get_children())
        if self.f2_entry_aluno.get() != '':
            self.f2_entry_aluno.insert(END, '%')
            nome = self.f2_entry_aluno.get()
            self.cursor.execute('''  
                            SELECT cod, 
                                nomes_alunos, 
                                series_alunos, 
                                nomes_livros, 
                                data_retirada, 
                                data_entregada 
                            FROM alunos WHERE nomes_alunos LIKE "%s"
                            ORDER BY nomes_alunos ASC
                                    ''' % nome)
            busca_nome = self.cursor.fetchall()
            for i in busca_nome:
                self.pesq_bd.insert("", END, values=i)
            self.limpa_f2()
            # Desconecta o BD
            self.desconecta_bd_f1()
        elif self.f2_entry_serie.get() != '':
            self.f2_entry_serie.insert(END, '%')
            serie = self.f2_entry_serie.get()
            self.cursor.execute('''  
                            SELECT cod, 
                                nomes_alunos, 
                                series_alunos, 
                                nomes_livros, 
                                data_retirada, 
                                data_entregada 
                            FROM alunos WHERE series_alunos LIKE "%s"
                            ORDER BY nomes_alunos ASC
                                    ''' % serie)
            busca_serie = self.cursor.fetchall()
            for i in busca_serie:
                self.pesq_bd.insert("", END, values=i)
            self.limpa_f2()
            # Desconecta o BD
            self.desconecta_bd_f1()
        elif self.f2_entry_livro.get() != '':
            self.f2_entry_livro.insert(END, '%')
            livro = self.f2_entry_livro.get()
            self.cursor.execute('''  
                            SELECT cod, 
                                nomes_alunos, 
                                series_alunos, 
                                nomes_livros, 
                                data_retirada, 
                                data_entregada 
                            FROM alunos WHERE nomes_livros LIKE "%s"
                            ORDER BY nomes_livros ASC
                                    ''' % livro)
            busca_livro = self.cursor.fetchall()
            for i in busca_livro:
                self.pesq_bd.insert("", END, values=i)
            self.limpa_f2()
            # Desconecta o BD
            self.desconecta_bd_f1()
        else:
            self.seleciona_f2()
            messagebox.showinfo(title='ERRO!',message='Insira algo para procurar!')
    def db_click(self, *args):
        # Funções de inicio
        self.limpar_f1()
        # Selecionar lista
        self.inser_bd.selection()
        # Função de duplo clique
        for n in self.inser_bd.selection():
            col1, col2, col3, col4, col5, col6 = self.inser_bd.item(n, 'values')
            self.f1_entry_cod.insert(END, col1)
            self.f1_entry_aluno.insert(END, col2)
            self.f1_entry_serie.insert(END, f'Você selecionou: {col3}')
            self.f1_variavel_serie.set(col3)
            self.f1_entry_livro.insert(END, col4)
            self.entry_data_retirada.insert(END, col5)
            self.entry_data_entrega.insert(END, col6)
    def db_click_t2(self, *args):
        # Funções de inicio
        self.limpar_f1()
        #entry
        self.variaveis_imp2()
        # Selecionar lista
        self.pesq_bd.selection()
        # Função de duplo clique
        for n in self.pesq_bd.selection():
            col1, col2, col3, col4, col5, col6 = self.pesq_bd.item(n, 'values')
            self.e_c.insert(END, col1)
            self.e_n.insert(END, col2)
            self.e_s.insert(END, col3)
            self.e_l.insert(END, col4)
            self.e_d1.insert(END, col5)
            self.e_d2.insert(END, col6)
            # Criação da tela de relatório
            consul_t2 = Tk()
            consul_t2.title(f'Aluno (a):  {(self.e_n.get())}')
            consul_t2.geometry('450x300')
            self.frame = Frame(consul_t2, bd=4, bg=self.cor1,
                            highlightbackground=self.cor2, highlightthickness=0.5, highlightcolor=self.cor2)
            self.frame.place(relx=0, rely=0, relwidth=1, relheight=1)
            # Criação das Labels
            # Label do titulo
            inf = Label(self.frame, text='Dados da coleta',
                            bg=self.cor1, fg=self.cor3,
                                    font=('Verdana', 14, 'bold'))
            inf.pack(padx=0,pady=0)
            # Label de Codigo
            c = Label(self.frame, text=f'Código de coleta:   {(self.e_c.get())}', anchor=W,
                    bg=self.cor1, fg=self.cor3,
                            font=('Verdana', 10, 'bold'))
            c.place(relx=0.02, rely=0.128, relheight=0.08, relwidth=1)
            # Label de Nome
            n = Label(self.frame, text=f'Nome do aluno:   {(self.e_n.get())}', anchor=W,
                            bg=self.cor1, fg=self.cor3,
                                    font=('Verdana', 10, 'bold'))
            n.place(relx=0.02, rely=0.256, relheight=0.08, relwidth=1)
            # Label de Serie
            s = Label(self.frame, text=f'Ano do aluno:   {(self.e_s.get())}', anchor=W,
                            bg=self.cor1, fg=self.cor3,
                                    font=('Verdana', 10, 'bold'))
            s.place(relx=0.02, rely=0.384, relheight=0.08, relwidth=1)
            # Label de Livro
            l = Label(self.frame, text=f'Nome do livro:   {(self.e_l.get())}', anchor=W,
                            bg=self.cor1, fg=self.cor3,
                                    font=('Verdana', 10, 'bold'))
            l.place(relx=0.02, rely=0.512, relheight=0.08, relwidth=1)
            # Label de Data de Retirada
            d1 = Label(self.frame, text=f'Data de retirada do livro:   {(self.e_d1.get())}', anchor=W,
                            bg=self.cor1, fg=self.cor3,
                                    font=('Verdana', 10, 'bold'))
            d1.place(relx=0.02, rely=0.64, relheight=0.08, relwidth=1)
            # Label de Data Entrega
            d2 = Label(self.frame, text=f'Data de entrega do livro:   {(self.e_d2.get())}', anchor=W,
                            bg=self.cor1, fg=self.cor3,
                                    font=('Verdana', 10, 'bold'))
            d2.place(relx=0.02, rely=0.768, relheight=0.08, relwidth=1)
            # Botão de imprimir relatório
            self.b_imp = Button(self.frame,
                                        bg=self.cor9, bd=0, activebackground=self.cor9, command=self.gerarelatorio,
                                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                        text='Imprimir', font=('Verdana', 10), cursor='hand2')
            self.b_imp.place(relx=0.02, rely=0.896, relheight=0.08, relwidth=0.46)
            # Botão de fechar janela
            self.b_close = Button(self.frame,
                                        bg=self.cor8, bd=0, activebackground=self.cor8, command=consul_t2.destroy,
                                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                        text='Fechar', font=('Verdana', 10), cursor='hand2')
            self.b_close.place(relx=0.52, rely=0.896, relheight=0.08, relwidth=0.46)
    def optionmenu(self, *args):
        # Limpar a Entry da Serie
        self.f1_entry_serie.delete(0, END)
        # Adicionar texto na Entry da Serie
        self.f1_entry_serie.insert(END, f'Você selecionou: {self.f1_variavel_serie.get()}')
        # Mudança de coloração da Entry
        if self.f1_variavel_serie != 'Selecionar':
            self.f1_drop_serie.config(bg=self.cor11, bd=0, highlightbackground=self.cor9,  relief=FLAT,
                                fg=self.cor7, font=('Verdana', 11), indicatoron='0', cursor='exchange', direction='flush')
    def cad_aluno(self):
        cor1 = '#B5B3A7'
        cad_tela = Tk()
        cad_tela.title('Cadastro de Alunos')
        cad_tela.geometry('700x470')
        cad_tela.attributes('-topmost', True)
        cad_tela.overrideredirect(False)
        Pmw.initialise(cad_tela)
        frame_cad = Frame(cad_tela, bd=0, bg=cor1,
                            highlightbackground=cor1, highlightthickness=0.5, highlightcolor=cor1)
        frame_cad.place(relx=0, rely=0, relwidth=1, relheight=1)

        ################# Criação das Labels #################

        # Label do titulo
        tit = Label(frame_cad, text='Dados de cadastro de alunos',
                        bg=cor1, fg=self.cor3,
                                font=('Verdana', 14, 'bold'))
        tit.pack(padx=0,pady=0)
        # Label de Nome
        nome = Label(frame_cad, text='Nome do Aluno', anchor=W,
                        bg=cor1, fg=self.cor3,
                        font=('Verdana', 9, 'bold'))
        nome.place(relx=0.02, rely=0.09, relheight=0.08, relwidth=1)
        # Label de Serie
        ano = Label(frame_cad, text='Ano do aluno', anchor=W,
                        bg=cor1, fg=self.cor3,
                        font=('Verdana', 9, 'bold'))
        ano.place(relx=0.02, rely=0.21, relheight=0.08, relwidth=1)
        # Label de Telefone
        tel = Label(frame_cad, text='Telefone de contato', anchor=W,
                        bg=cor1, fg=self.cor3,
                        font=('Verdana', 9, 'bold'))
        tel.place(relx=0.02, rely=0.33, relheight=0.08, relwidth=1)
        # Label de Identificação
        ra = Label(frame_cad, text='Identificação do Aluno', anchor=W,
                        bg=cor1, fg=self.cor3,
                        font=('Verdana', 9, 'bold'))
        ra.place(relx=0.02, rely=0.45, relheight=0.08, relwidth=1)
        # Label de E-mail
        mail = Label(frame_cad, text='E-mail de contato', anchor=W,
                        bg=cor1, fg=self.cor3,
                        font=('Verdana', 9, 'bold'))
        mail.place(relx=0.02, rely=0.57, relheight=0.08, relwidth=1)
        # Label de Mãe
        mae = Label(frame_cad, text='Nome da Mãe ou Responsável', anchor=W,
                        bg=cor1, fg=self.cor3,
                        font=('Verdana', 9, 'bold'))
        mae.place(relx=0.02, rely=0.69, relheight=0.08, relwidth=1)
        # Label de Pai
        pai = Label(frame_cad, text='Nome do Pai ou Responsável', anchor=W,
                        bg=cor1, fg=self.cor3,
                        font=('Verdana', 9, 'bold'))
        pai.place(relx=0.02, rely=0.81, relheight=0.08, relwidth=1)

        ################# Criação das Entrys #################

        # Entry do Codigo do Aluno
        entry_cod = Entry(frame_cad)
        # Entry do Nome do Aluno selecionado
        entry_nome = Entry(frame_cad,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10,fg=self.cor2, insertontime='0')
        entry_nome.place(relx=0.02, rely=0.16, relheight=0.06, relwidth=0.3)
        # Entry do Ano do Aluno selecionado
        entry_ano = Entry(frame_cad,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10,fg=self.cor2, insertontime='0')
        entry_ano.place(relx=0.02, rely=0.28, relheight=0.06, relwidth=0.3)
        # Entry do Telefone do Aluno selecionado
        entry_fone = Entry(frame_cad,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10,fg=self.cor2, insertontime='0')
        entry_fone.place(relx=0.02, rely=0.4, relheight=0.06, relwidth=0.3)
        # Entry da Identificação do Aluno selecionado
        entry_cpf = Entry(frame_cad,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10,fg=self.cor2, insertontime='0')
        entry_cpf.place(relx=0.02, rely=0.52, relheight=0.06, relwidth=0.3)
        # Entry de E-mail do Aluno selecionado
        entry_email = Entry(frame_cad,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10,fg=self.cor2, insertontime='0')
        entry_email.place(relx=0.02, rely=0.64, relheight=0.06, relwidth=0.3)
        # Entry de E-mail do Aluno selecionado
        entry_mae = Entry(frame_cad,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10,fg=self.cor2, insertontime='0')
        entry_mae.place(relx=0.02, rely=0.76, relheight=0.06, relwidth=0.3)
        # Entry de E-mail do Aluno selecionado
        entry_pai = Entry(frame_cad,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10,fg=self.cor2, insertontime='0')
        entry_pai.place(relx=0.02, rely=0.88, relheight=0.06, relwidth=0.3)

        ################# Criação dos Comandos #################

        def seleciona():
            #Selecionar a Treeview
            cad_bd.delete(*cad_bd.get_children())
            # Conectar o BD
            self.conecta_bd_f1()
            # Criar variavel sem o self
            lista_cad = self.cursor.execute(''' 
                            SELECT cod_1, 
                                nome, 
                                serie, 
                                telefone, 
                                identificacao, 
                                email,
                                mae,
                                pai
                            FROM identidade ORDER BY nome ASC; 
                                    ''')
            # Adicionar na Treeview
            for i in lista_cad:
                cad_bd.insert("", END, values=i)
            # Desconectar o BD
            self.desconecta_bd_f1()
        def variaveis():
            self.cod_get=entry_cod.get()
            self.n_get_cad=entry_nome.get()
            self.a_get_cad=entry_ano.get()
            self.f_get_cad=entry_fone.get()
            self.c_get_cad=entry_cpf.get()
            self.e_get_cad=entry_email.get()
            self.mae_get_cad=entry_mae.get()
            self.pai_get_cad=entry_pai.get()
        def variaveis1():
            entry_ano.insert(END, '° Ano')
            self.cod_get=entry_cod.get()
            self.n_get_cad=entry_nome.get()
            self.a_get_cad=entry_ano.get()
            self.f_get_cad=entry_fone.get()
            self.c_get_cad=entry_cpf.get()
            self.e_get_cad=entry_email.get()
            self.mae_get_cad=entry_mae.get()
            self.pai_get_cad=entry_pai.get()
        def clean():
            entry_cod.delete(0, END)
            entry_nome.delete(0, END)
            entry_ano.delete(0, END)
            entry_fone.delete(0, END)
            entry_cpf.delete(0, END)
            entry_email.delete(0, END)
            entry_mae.delete(0, END)
            entry_pai.delete(0, END)
        def bt__add():
            variaveis1()
            
            if entry_nome.get()=='' or entry_ano.get()=='' or entry_fone.get()=='' or entry_cpf.get()=='':
                messagebox.showinfo(title='ERRO!', message='Insira todos os dados!')
                return
            self.conecta_bd_f1()
            # Função de colocar no BD
            
            self.cursor.execute('''
                    INSERT INTO identidade 
                        (nome, 
                        serie, 
                        telefone, 
                        identificacao, 
                        email,
                        mae,
                        pai)
                    VALUES (?,?,?,?,?,?,?)''', 
                        (self.n_get_cad, 
                        self.a_get_cad, 
                        self.f_get_cad, 
                        self.c_get_cad,
                        self.e_get_cad,
                        self.mae_get_cad,
                        self.pai_get_cad
                        ))
            self.conn.commit()
            self.desconecta_bd_f1()
            seleciona()
            clean()
        def bt__del():
            # Chamar Variaveis
            variaveis()
            # Conectar o BD
            self.conecta_bd_f1()
            # Função de deletar do BD
            try:
                self.cursor.execute('''DELETE FROM identidade WHERE cod_1 = {} '''.format(self.cod_get))
                self.conn.commit()
            except:
                messagebox.showinfo(title='ERRO!',message='Selecione com 2 cliques para ser apagado!')
            # Desconectar o BD
            self.desconecta_bd_f1()
            # Limpar as Entrys
            clean()
            # Atualizar a lista
            seleciona()
            # VOltar Bt de ADD
            # Button de Adicionar informações no BD
            bt_add = Button(frame_cad, command=bt__add,
                        bg=self.cor9, bd=0, activebackground=self.cor11,
                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                        text='Cadastrar', font=('Verdana', 10), cursor='hand2')
            bt_add.place(relx=0.37, rely=0.88, relheight=0.06, relwidth=0.14)
        def bt__alt():
            # Variaveis
            variaveis()
            # Conectar ao BD
            self.conecta_bd_f1()
            # Código de alterar os dados
            self.cursor.execute('''
                        UPDATE identidade 
                            SET nome = ?, 
                            serie = ?, 
                            telefone = ?, 
                            identificacao = ?, 
                            email = ?,
                            mae = ?,
                            pai = ? 
                        WHERE cod_1 = ? 
                        ''', 
                        (self.n_get_cad, 
                        self.a_get_cad, 
                        self.f_get_cad, 
                        self.c_get_cad,
                        self.e_get_cad,
                        self.mae_get_cad,
                        self.pai_get_cad,
                        self.cod_get))
            self.conn.commit()
            # Desconectar o BD
            self.desconecta_bd_f1()
            # Atualizar a Tabela
            seleciona()
            # Limpar as Entrys
            clean()
            # Voltar BT de Add
            # Button de Adicionar informações no BD
            bt_add = Button(frame_cad, command=bt__add,
                        bg=self.cor9, bd=0, activebackground=self.cor11,
                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                        text='Cadastrar', font=('Verdana', 10), cursor='hand2')
            bt_add.place(relx=0.37, rely=0.88, relheight=0.06, relwidth=0.14)
        def db_click(*args):
            # Funções de inicio
            clean()
            #entry
            variaveis()
            # Selecionar lista
            cad_bd.selection()
            # Função de duplo clique
            for n in cad_bd.selection():
                col1, col2, col3, col4, col5, col6, col7, col8 = cad_bd.item(n, 'values')
                entry_cod.insert(END, col1)
                entry_nome.insert(END, col2)
                entry_ano.insert(END, col3)
                entry_fone.insert(END, col4)
                entry_cpf.insert(END, col5)
                entry_email.insert(END, col6)
                entry_mae.insert(END, col7)
                entry_pai.insert(END, col8)
                # Criação do Botão de Informações
                bt_inf = Button(frame_cad, command=bt__inf,
                        bg=self.cor9, bd=0, activebackground=self.cor11,
                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                        text='Informações', font=('Verdana', 10), cursor='hand2')
                bt_inf.place(relx=0.37, rely=0.88, relheight=0.06, relwidth=0.14)
        def bt__inf():
            # Criação da tela de relatório
            consul_t2 = Tk()
            consul_t2.title('Informações do aluno')
            consul_t2.attributes('-topmost', True)
            consul_t2.geometry('550x500')
            frame = Frame(consul_t2, bd=4, bg=self.cor1,
                            highlightbackground=self.cor2, highlightthickness=0.5, highlightcolor=self.cor2)
            frame.place(relx=0, rely=0, relwidth=1, relheight=1)
            # Criação das Labels                # Label do titulo
            inf = Label(frame, text='Dados da coleta',
                                bg=self.cor1, fg=self.cor3,
                                        font=('Verdana', 14, 'bold'))
            inf.pack(padx=0,pady=0)
            # Label de Codigo
            c = Label(frame, text=f'Código de coleta:   {(entry_cod.get())}', anchor=W,
                        bg=self.cor1, fg=self.cor3,
                                font=('Verdana', 10, 'bold'))
            c.place(relx=0.02, rely=0.1, relheight=0.08, relwidth=1)
            # Label de Nome
            n = Label(frame, text=f'Nome do aluno:   {(entry_nome.get())}', anchor=W,
                                bg=self.cor1, fg=self.cor3,
                                        font=('Verdana', 10, 'bold'))
            n.place(relx=0.02, rely=0.2, relheight=0.08, relwidth=1)
            # Label de Serie
            s = Label(frame, text=f'Ano do aluno:   {(entry_ano.get())}', anchor=W,
                                bg=self.cor1, fg=self.cor3,
                                        font=('Verdana', 10, 'bold'))
            s.place(relx=0.02, rely=0.3, relheight=0.08, relwidth=1)
            # Label de Livro
            f = Label(frame, text=f'Telefone de contato:   {(entry_fone.get())}', anchor=W,
                                bg=self.cor1, fg=self.cor3,
                                        font=('Verdana', 10, 'bold'))
            f.place(relx=0.02, rely=0.4, relheight=0.08, relwidth=1)
            # Label de Data de Retirada
            cp = Label(frame, text=f'Documento do Aluno:   {(entry_cpf.get())}', anchor=W,
                                bg=self.cor1, fg=self.cor3,
                                        font=('Verdana', 10, 'bold'))
            cp.place(relx=0.02, rely=0.5, relheight=0.08, relwidth=1)
            # Label de Data Entrega
            em = Label(frame, text=f'E-mail de contato:   {(entry_email.get())}', anchor=W,
                                bg=self.cor1, fg=self.cor3,
                                        font=('Verdana', 10, 'bold'))
            em.place(relx=0.02, rely=0.6, relheight=0.08, relwidth=1)
            # Label de Data Entrega
            nm = Label(frame, text=f'Nome da Mãe:   {(entry_mae.get())}', anchor=W,
                                bg=self.cor1, fg=self.cor3,
                                        font=('Verdana', 10, 'bold'))
            nm.place(relx=0.02, rely=0.7, relheight=0.08, relwidth=1)
            # Label de Data Entrega
            np = Label(frame, text=f'Nome do Pai:   {(entry_pai.get())}', anchor=W,
                                bg=self.cor1, fg=self.cor3,
                                        font=('Verdana', 10, 'bold'))
            np.place(relx=0.02, rely=0.8, relheight=0.08, relwidth=1)
            # Destruir Bt de Inf
            bt_add = Button(frame_cad, command=bt__add,
                        bg=self.cor9, bd=0, activebackground=self.cor11,
                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                        text='Cadastrar', font=('Verdana', 10), cursor='hand2')
            bt_add.place(relx=0.37, rely=0.88, relheight=0.06, relwidth=0.14)
            clean()

        ################# Criação dos Button #################

        # Button de Adicionar informações no BD 
        bt_add = Button(frame_cad, command=bt__add,
                        bg=self.cor9, bd=0, activebackground=self.cor11,
                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                        text='Cadastrar', font=('Verdana', 10), cursor='hand2')
        bt_add.place(relx=0.37, rely=0.88, relheight=0.06, relwidth=0.14)
        # Button de Limpar informações nas Entrys
        bt_cle = Button(frame_cad, command=clean,
                        bg=self.cor6, bd=0, activebackground=self.cor11,
                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                        text='Limpar', font=('Verdana', 10), cursor='hand2')
        bt_cle.place(relx=0.524, rely=0.88, relheight=0.06, relwidth=0.14)
        # Button de Alterar informações no BD
        bt_alt = Button(frame_cad, command=bt__alt,
                        bg=self.cor6, bd=0, activebackground=self.cor11,
                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                        text='Alterar', font=('Verdana', 10), cursor='hand2')
        bt_alt.place(relx=0.676, rely=0.88, relheight=0.06, relwidth=0.14)
        # Button de Deletar informações no BD
        bt_del = Button(frame_cad, command=bt__del,
                        bg=self.cor8, bd=0, activebackground=self.cor11,
                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                        text='Excluir', font=('Verdana', 10), cursor='hand2')
        bt_del.place(relx=0.83, rely=0.88, relheight=0.06, relwidth=0.14)
        
        ################# Criação da Treeview #################
        
        # Criação da Treeview
        cad_bd = ttk.Treeview(frame_cad, columns=('col1', 'col2', 'col3', 'col4', 'col5', 'col6', 'col7', 'col8'))
        cad_bd.place(relx=0.37, rely=0.1, relheight=0.76, relwidth=0.6)
        # Cabeçalho da Treeview
        cad_bd.heading('#0', text='')
        cad_bd.heading('#1', text='#')
        cad_bd.heading('#2', text='Nome')
        cad_bd.heading('#3', text='Ano')
        cad_bd.heading('#4', text='Telefone')
        cad_bd.heading('#5', text='Documento')
        cad_bd.heading('#6', text='E-mail')
        cad_bd.heading('#7', text='Nome da Mãe')
        cad_bd.heading('#8', text='Nome do Pai')
        # Tamanho das colunas
        cad_bd.column('#0', width=0,stretch=0)
        cad_bd.column('#1', width=1,stretch=1)
        cad_bd.column('#2', width=75,stretch=1)
        cad_bd.column('#3', width=20,stretch=1)
        cad_bd.column('#4', width=85,stretch=1)
        cad_bd.column('#5', width=70,stretch=1)
        cad_bd.column('#6', width=90,stretch=1)
        cad_bd.column('#7', width=0,stretch=0)
        cad_bd.column('#8', width=0,stretch=0)
        cad_bd.bind('<Double-1>', db_click)
        seleciona()
        # MSG BOX
        # Variaveis Botões
        msg_bt_add = 'Adicionar Aluno'
        msg_bt_cle = 'Limpar as as Entradas de Texto'
        msg_bt_alt = 'Alterar dados do Aluno'
        msg_bt_del = 'Excluir Aluno'
        # Variaveis Botões
        msg_et_nome = 'Digite o nome do Aluno a ser cadastrado'
        msg_et_ano = 'Selecione o Ano do aluno'
        msg_et_telefone = 'Digite o Telefone de Contato'
        msg_et_doc = 'Digite uma documentação do Aluno'
        msg_et_email = 'Digite um E-mail de contato'
        msg_et_mae = 'Nome da Mãe ou Responsavel pelo aluno'
        msg_et_pai = 'Nome do Pai ou Responsavel pelo aluno'

        # msg Button Adicionar
        box_add = Pmw.Balloon(frame_cad)
        box_add.bind(bt_add, msg_bt_add)
        # msg Button Limpar
        box_cle = Pmw.Balloon(frame_cad)
        box_cle.bind(bt_cle, msg_bt_cle)
        # msg Button Alterar
        box_alt = Pmw.Balloon(frame_cad)
        box_alt.bind(bt_alt, msg_bt_alt)
        # msg Button Excluir
        box_del = Pmw.Balloon(frame_cad)
        box_del.bind(bt_del, msg_bt_del)
    def cad_livro(self):
        cor1 = '#B5B3A7'
        cad_tela = Tk()
        cad_tela.title('Cadastro de Livros')
        cad_tela.geometry('700x470')
        cad_tela.attributes('-topmost', True)
        cad_tela.overrideredirect(False)
        frame_cad = Frame(cad_tela, bd=0, bg=cor1,
                            highlightbackground=cor1, highlightthickness=0.5, highlightcolor=cor1)
        frame_cad.place(relx=0, rely=0, relwidth=1, relheight=1)

        ################# Criação das Labels #################

        # Label do titulo
        tit = Label(frame_cad, text='Dados de cadastro de Livros',
                        bg=cor1, fg=self.cor3,
                                font=('Verdana', 14, 'bold'))
        tit.pack(padx=0,pady=0)
        # Label de Nome
        nome = Label(frame_cad, text='Nome do Livro', anchor=W,
                        bg=cor1, fg=self.cor3,
                        font=('Verdana', 9, 'bold'))
        nome.place(relx=0.02, rely=0.128, relheight=0.08, relwidth=1)
        # Label de Serie
        ano = Label(frame_cad, text='Ano de Lançamento do Livro', anchor=W,
                        bg=cor1, fg=self.cor3,
                        font=('Verdana', 9, 'bold'))
        ano.place(relx=0.02, rely=0.256, relheight=0.08, relwidth=1)
        # Label de Autor
        aut = Label(frame_cad, text='Autor do Livro', anchor=W,
                        bg=cor1, fg=self.cor3,
                        font=('Verdana', 9, 'bold'))
        aut.place(relx=0.02, rely=0.384, relheight=0.08, relwidth=1)
        # Label de Genero
        gen = Label(frame_cad, text='Genero do Livro', anchor=W,
                        bg=cor1, fg=self.cor3,
                        font=('Verdana', 9, 'bold'))
        gen.place(relx=0.02, rely=0.512, relheight=0.08, relwidth=1)
        # Label de Páginas do Livro
        gen = Label(frame_cad, text='N° de páginas do Livro', anchor=W,
                        bg=cor1, fg=self.cor3,
                        font=('Verdana', 9, 'bold'))
        gen.place(relx=0.02, rely=0.64, relheight=0.08, relwidth=1)

        ################# Criação das Entrys #################

        entry_cod = Entry(frame_cad)
        # Entry do Nome do Livro
        entry_nome = Entry(frame_cad,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10,fg=self.cor2, insertontime='0')
        entry_nome.place(relx=0.02, rely=0.198, relheight=0.06, relwidth=0.3)
        # Entry do Ano do Livro
        entry_ano = Entry(frame_cad,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10,fg=self.cor2, insertontime='0')
        entry_ano.place(relx=0.02, rely=0.326, relheight=0.06, relwidth=0.3)
        # Entry do Autor do Livro
        entry_autor = Entry(frame_cad,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10,fg=self.cor2, insertontime='0')
        entry_autor.place(relx=0.02, rely=0.454, relheight=0.06, relwidth=0.3)
        # Entry do Genero do Livro
        entry_genero = Entry(frame_cad,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10,fg=self.cor2, insertontime='0')
        entry_genero.place(relx=0.02, rely=0.582, relheight=0.06, relwidth=0.3)
        # Entry do Número de páginas do livro
        entry_numero = Entry(frame_cad,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10,fg=self.cor2, insertontime='0')
        entry_numero.place(relx=0.02, rely=0.71, relheight=0.06, relwidth=0.3)

        ################# Criação dos Comandos #################

        def seleciona():
            #Selecionar a Treeview
            cad_bd.delete(*cad_bd.get_children())
            # Conectar o BD
            self.conecta_bd_f1()
            # Criar variavel sem o self
            lista_cad = self.cursor.execute(''' 
                            SELECT cod_2, 
                                nome_l, 
                                ano, 
                                autor, 
                                genero,
                                numero 
                            FROM livros ORDER BY nome_l ASC; 
                                    ''')
            # Adicionar na Treeview
            for i in lista_cad:
                cad_bd.insert("", END, values=i)
            # Desconectar o BD
            self.desconecta_bd_f1()
        def variaveis():
            self.cod_get_l=entry_cod.get()
            self.n_get_cad_l=entry_nome.get()
            self.a_get_cad_l=entry_ano.get()
            self.at_get_cad_l=entry_autor.get()
            self.g_get_cad_l=entry_genero.get()
            self.p_get_cad_l=entry_numero.get()
        def clean():
            entry_nome.delete(0, END)
            entry_ano.delete(0, END)
            entry_autor.delete(0, END)
            entry_genero.delete(0, END)
            entry_numero.delete(0, END)
        def bt__add():
            variaveis()
            if entry_nome.get()=='':
                messagebox.showinfo(title='ERRO!', message='Insira todos os dados!')
                return
            self.conecta_bd_f1()
            # Função de colocar no BD
            self.cursor.execute('''
                    INSERT INTO livros 
                        (nome_l, 
                        ano, 
                        autor, 
                        genero,
                        numero
                        )
                    VALUES (?,?,?,?,?)''', 
                        (self.n_get_cad_l, 
                        self.a_get_cad_l, 
                        self.at_get_cad_l, 
                        self.g_get_cad_l,
                        self.p_get_cad_l
                        ))
            self.conn.commit()
            self.desconecta_bd_f1()
            seleciona()
            clean()
        def bt__del():
            # Chamar Variaveis
            variaveis()
            # Conectar o BD
            self.conecta_bd_f1()
            # Função de deletar do BD
            try:
                self.cursor.execute('''DELETE FROM livros WHERE cod_2 = {} '''.format(self.cod_get_l))
                self.conn.commit()
            except:
                messagebox.showinfo(title='ERRO!',message='Selecione com 2 cliques para ser apagado!')
            # Desconectar o BD
            self.desconecta_bd_f1()
            # Limpar as Entrys
            clean()
            # Atualizar a lista
            seleciona()
        def bt__alt():
            # Variaveis
            variaveis()
            # Conectar ao BD
            self.conecta_bd_f1()
            # Código de alterar os dados
            self.cursor.execute('''
                        UPDATE livros
                            SET nome_l = ?, 
                            ano = ?, 
                            autor = ?, 
                            genero = ?, 
                            numero = ?
                        WHERE cod_2 = ? 
                        ''', 
                        (self.n_get_cad_l, 
                        self.a_get_cad_l, 
                        self.at_get_cad_l, 
                        self.g_get_cad_l,
                        self.p_get_cad_l,
                        self.cod_get_l))
            self.conn.commit()
            # Desconectar o BD
            self.desconecta_bd_f1()
            # Atualizar a Tabela
            seleciona()
            # Limpar as Entrys
            clean()
        def db_click(*args):
            # Funções de inicio
            clean()
            #entry
            variaveis()
            # Selecionar lista
            cad_bd.selection()
            # Função de duplo clique
            for n in cad_bd.selection():
                col1, col2, col3, col4, col5, col6 = cad_bd.item(n, 'values')
                entry_cod.insert(END, col1)
                entry_nome.insert(END, col2)
                entry_ano.insert(END, col3)
                entry_autor.insert(END, col4)
                entry_genero.insert(END, col5)
                entry_numero.insert(END, col6)
        
        ################# Criação dos Button #################

        # Button de Adicionar informações no BD
        bt_add = Button(frame_cad, command=bt__add,
                        bg=self.cor9, bd=0, activebackground=self.cor11,
                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                        text='Cadastrar', font=('Verdana', 10), cursor='hand2')
        bt_add.place(relx=0.02, rely=0.82, relheight=0.06, relwidth=0.14)
        # Button de Limpar informações nas Entrys
        bt_cle = Button(frame_cad, command=clean,
                        bg=self.cor6, bd=0, activebackground=self.cor11,
                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                        text='Limpar', font=('Verdana', 10), cursor='hand2')
        bt_cle.place(relx=0.18, rely=0.82, relheight=0.06, relwidth=0.14)
        # Button de Deletar informações no BD
        bt_del = Button(frame_cad, command=bt__del,
                        bg=self.cor8, bd=0, activebackground=self.cor11,
                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                        text='Excluir', font=('Verdana', 10), cursor='hand2')
        bt_del.place(relx=0.18, rely=0.9, relheight=0.06, relwidth=0.14)
        # Button de Alterar informações no BD
        bt_alt = Button(frame_cad, command=bt__alt,
                        bg=self.cor6, bd=0, activebackground=self.cor11,
                        highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                        text='Alterar', font=('Verdana', 10), cursor='hand2')
        bt_alt.place(relx=0.02, rely=0.9, relheight=0.06, relwidth=0.14)

        ################# Criação da Treeview #################
        
        # Criação da Treeview
        cad_bd = ttk.Treeview(frame_cad, columns=('col1', 'col2', 'col3', 'col4', 'col5', 'col6'))
        cad_bd.place(relx=0.37, rely=0.198, relheight=0.76, relwidth=0.6)
        # Cabeçalho da Treeview
        cad_bd.heading('#0', text='')
        cad_bd.heading('#1', text='#')
        cad_bd.heading('#2', text='Nome')
        cad_bd.heading('#3', text='Ano')
        cad_bd.heading('#4', text='Autor')
        cad_bd.heading('#5', text='Genero')
        cad_bd.heading('#6', text='N°')
        # Tamanho das colunas
        cad_bd.column('#0', width=0,stretch=0)
        cad_bd.column('#1', width=1,stretch=1)
        cad_bd.column('#2', width=100,stretch=1)
        cad_bd.column('#3', width=15,stretch=1)
        cad_bd.column('#4', width=105,stretch=1)
        cad_bd.column('#5', width=75,stretch=1)
        cad_bd.column('#6', width=15,stretch=1)
        cad_bd.bind('<Double-1>', db_click)
        seleciona()
    def cad_escolas(self):
        print()
    def cad_usuarios(self):
        print()

class Application(Funcs,pdf,bd):
    def __init__(self):
        self.root = root
        self.cores()
        self.tela1()
        self.frames_abas()
        self.cadastro()
        self.aba_login()
        self.date_today()
        self.labels()
        self.botoes()
        self.entry()
        self.treeview_f1()
        self.treeview_f2()
        self.montatabelas_f1()
        self.atrasados()
        self.seleciona_f1()
        self.seleciona_f2()
        self.limpar_f1()
        self.msg_box()
        self.anotacoes()
        
        #self.menus()
        root.mainloop()
    def cores(self):
        self.cor1 = '#002B36' # Cor de fundo
        self.cor2 = '#00252e' # Cor de borda
        self.cor3 = '#ffffff' # Branco
        self.cor4 = '#073642' # Escuro
        self.cor5 = '#A9BDBD' # Claro
        self.cor6 = '#3f98d7' # info
        self.cor7 = '#060606' # Quase preto
        self.cor8 = '#d95092' # Danger
        self.cor9 = '#44aca4' # Success
        self.cor10 = '#d05e2f' # Warning
        self.cor11 = '#FAF7E7' # Fundo da Entry
    def tela1(self):
        # Janela principal do aplicativo
        self.root.title('Biblioteca')
        self.root.geometry('450x400')
        self.root.iconbitmap('ARQUIVOS\IMAGENS\icone.ico')
        #self.root.resizable(False, False)
        #self.root.minsize(width=360, height=310)
        #self.root.maxsize(width=660, height=610)
    def aba_login(self):
        # Label de Login na parte superior
        self.l_label_tit = Label(self.frame_tela0,text='LOGIN',bg=self.cor1,fg=self.cor3,
                                font=('Helvetica',30))
        self.l_label_tit.place(relx=0.02,rely=0.001)
        # Linha de separação do Login na parte superior
        self.l_linha_tit = Label(self.frame_tela0,text='',bg=self.cor9,
                                font=('Helvetica',1))
        self.l_linha_tit.place(relx=0.02,rely=0.12,relwidth=0.96)
        # Label de nome de usuario
        self.l_label_usuario = Label(self.frame_tela0,text='Nome',bg=self.cor1,fg=self.cor3,
                                font=('Helvetica',13))
        self.l_label_usuario.place(relx=0.06,rely=0.2)
        # Entry de usuario
        self.l_entry_usu = Entry(self.frame_tela0,
                                    bg=self.cor11,bd=3,highlightbackground=self.cor9,relief=FLAT,
                                    highlightthickness=2,highlightcolor=self.cor9,fg=self.cor2,insertontime='700')
        self.l_entry_usu.place(relx=0.02,rely=0.27,relheight=0.07,relwidth=0.6)
        # Label de senha do usuario
        self.l_label_senha = Label(self.frame_tela0,text='Senha',bg=self.cor1,fg=self.cor3,
                                font=('Helvetica',13))
        self.l_label_senha.place(relx=0.06,rely=0.4)
        # Comando senha
        def valida_usuario():
            if self.l_entry_sen1.get() != 'admin' and self.l_entry_usu.get() != 'admin':
                messagebox.showinfo(title='ERRO!',message='Credenciais invalidas')
            else:
                self.frame_tela0.destroy()
        # Entry de senha
        self.l_entry_sen1 = Entry(self.frame_tela0,show='*',
                                    bg=self.cor11,bd=3,highlightbackground=self.cor9,relief=FLAT,
                                    highlightthickness=2,highlightcolor=self.cor9,fg=self.cor2,insertontime='700')
        self.l_entry_sen1.place(relx=0.02,rely=0.47,relheight=0.07,relwidth=0.6)
        # Comando Checkbox
        c_v1=IntVar(value=0)
        def show_pass():
            if (c_v1.get() == 1):
                self.l_entry_sen1.config(show='')
            else:
                self.l_entry_sen1.config(show='*')
        # Checkbox e label de mostrar a senha
        self.l_ck_senha = Checkbutton(self.frame_tela0,bg=self.cor1,activebackground=self.cor1, bd=3,
                                    command=show_pass,variable=c_v1)
        self.l_ck_senha.place(relx=0.02, rely=0.55)
        self.l_label_ck = Label(self.frame_tela0,text='Mostrar senha',bg=self.cor1,fg=self.cor3,
                                font=('Helvetica', 10))
        self.l_label_ck.place(relx=0.07,rely=0.56)
        # Botão de Login
        self.l_bt_login = Button(self.frame_tela0,
                                    bg=self.cor9,bd=0,activebackground=self.cor11,command=valida_usuario,
                                    highlightbackground=self.cor2,highlightthickness=1,fg=self.cor3,
                                    text='Entrar',font=('Verdana', 12),cursor='hand2')
        self.l_bt_login.place(relx=0.05, rely=0.63, relheight=0.08, relwidth=0.2)
        # Botão de Cancelar
        self.l_bt_can = Button(self.frame_tela0,
                                    bg=self.cor8,bd=0,activebackground=self.cor8,command=root.destroy,
                                    highlightbackground=self.cor2,highlightthickness=1,fg=self.cor3,
                                    text='Cancelar',font=('Verdana', 12),cursor='hand2')
        self.l_bt_can.place(relx=0.29, rely=0.63, relheight=0.08, relwidth=0.2)
        # Função de não existe
        def nao_existe():
            messagebox.showinfo(title='Erro', message='Essa função infelizmente ainda não existe')
        # Botão de Esqueceu a senha
        self.l_bt_e_senha = Button(self.frame_tela0,
                                    bg=self.cor1,bd=0,activebackground=self.cor1,command=nao_existe,
                                    highlightbackground=self.cor2,highlightthickness=1,fg=self.cor9,activeforeground=self.cor9,
                                    text='Esqueceu a senha?',font=('Verdana',8),cursor='hand2')
        self.l_bt_e_senha.place(relx=0.01, rely=0.82)
        # Botão de não tem usuario
        self.l_bt_n_lo = Button(self.frame_tela0,
                                    bg=self.cor1,bd=0,activebackground=self.cor1,command=nao_existe,
                                    highlightbackground=self.cor2,highlightthickness=1,fg=self.cor9,activeforeground=self.cor9,
                                    text='Ainda não tem cadastro?',font=('Verdana',8),cursor='hand2')
        self.l_bt_n_lo.place(relx=0.01, rely=0.9)
    def labels(self):

        ################### Tela 1 Inicio ###################

        # Label aluno tela 1
        self.f1_label_aluno = Label(self.frame_tela1, text='Aluno', bg=self.cor1, fg=self.cor3,
                                font=('Verdana', 10, 'bold'))
        self.f1_label_aluno.place(relx=0.02, rely=0.001, relheight=0.08, relwidth=0.17)
        # Label serie tela 1
        self.f1_label_serie = Label(self.frame_tela1, text='Serie do aluno', bg=self.cor1, fg=self.cor3,
                                font=('Verdana', 10, 'bold'))
        self.f1_label_serie.place(relx=0.65, rely=0.001, relheight=0.08, relwidth=0.30)
        # Label livro tela 1
        self.f1_label_livro = Label(self.frame_tela1, text='Livro', bg=self.cor1, fg=self.cor3,
                                font=('Verdana', 10, 'bold'))
        self.f1_label_livro.place(relx=0.02, rely=0.183, relheight=0.08, relwidth=0.17)
        # Label Cadastros
        self.f1_label_cadastro = Label(self.frame_tela1, text='Cadastros', bg=self.cor1, fg=self.cor3,
                                font=('Verdana', 10, 'bold'))
        self.f1_label_cadastro.place(relx=0.02, rely=0.35, relheight=0.08, relwidth=0.25)

        ################### Tela 2 Consulta ###################

        # Label pesq aluno
        self.f2_l_aluno = Label(self.frame_tela2, text='Aluno', bg=self.cor1, fg=self.cor3,
                                font=('Verdana', 10, 'bold'))
        self.f2_l_aluno.place(relx=0.014, rely=0.002, relheight=0.08, relwidth=0.17)
        # Label pesq serie
        self.f2_l_serie = Label(self.frame_tela2, text='Ano', bg=self.cor1, fg=self.cor3,
                                font=('Verdana', 10, 'bold'))
        self.f2_l_serie.place(relx=0.27, rely=0.002, relheight=0.08, relwidth=0.17)
        # Label pesq livro
        self.f2_l_livro = Label(self.frame_tela2, text='Livro', bg=self.cor1, fg=self.cor3,
                                font=('Verdana', 10, 'bold'))
        self.f2_l_livro.place(relx=0.528, rely=0.002, relheight=0.08, relwidth=0.17)
    def frames_abas(self):
        # Personalizando as tabs
        abas_style = ttk.Style()
        abas_style.theme_use('default') # clam , default , alt , vista
        abas_style.configure('TNotebook', background=self.cor1)
        abas_style.configure('TNotebook.Tab', background=self.cor4, foreground=self.cor3)
        abas_style.map("TNotebook.Tab", background= [("selected", self.cor1)])
        # Abas
        self.abas = ttk.Notebook(self.root)
        self.abas.place(relx=0, rely=0, relwidth=1, relheight=1)
        # Tela 1
        self.aba1 = Frame(self.abas)
        self.aba1.configure(background=self.cor1, bd=10)
        self.abas.add(self.aba1, text='       Inicio      ')
        # Tela 2
        self.aba2 = Frame(self.abas)
        self.aba2.configure(background=self.cor1, bd=10)
        self.abas.add(self.aba2, text='     Consulta     ')
        # Tela 3
        self.aba3 = Frame(self.abas)
        self.aba3.configure(background=self.cor1, bd=10)
        self.abas.add(self.aba3, text='     Cadastro     ')
        # Tela 4
        self.aba4 = Frame(self.abas)
        self.aba4.configure(background=self.cor1, bd=10)
        self.abas.add(self.aba4, text='     Anotações     ')
        # Tela 5
        self.aba5 = Frame(self.abas)
        self.aba5.configure(background=self.cor1, bd=10)
        self.abas.add(self.aba5, text='     Atrasados     ')
        # Frame da Tela de Retirada de livro
        self.frame_tela1 = Frame(self.aba1, bd=4, bg=self.cor1,
                                highlightbackground=self.cor2, highlightthickness=0.5, highlightcolor=self.cor2)
        self.frame_tela1.place(relx=0, rely=0, relwidth=1, relheight=1)
        # Frame da Tela de Pesquisa
        self.frame_tela2 = Frame(self.aba2, bd=4, bg=self.cor1,
                                highlightbackground=self.cor2, highlightthickness=0.5, highlightcolor=self.cor2)
        self.frame_tela2.place(relx=0, rely=0, relwidth=1, relheight=1)
        # Frame da Tela Cadastros
        self.frame_tela3 = Frame(self.aba3, bd=4, bg=self.cor11,
                                highlightbackground=self.cor2, highlightthickness=0.5, highlightcolor=self.cor2)
        self.frame_tela3.place(relx=0, rely=0, relwidth=1, relheight=1)
        # Frame da Tela Atrasados
        self.frame_tela5 = Frame(self.aba5, bd=4, bg=self.cor11,
                                highlightbackground=self.cor2, highlightthickness=0.5, highlightcolor=self.cor2)
        self.frame_tela5.place(relx=0, rely=0, relwidth=1, relheight=1)
        # Frame da Tela de login
        self.frame_tela0 = Frame(root, bd=4, bg=self.cor1,
                                highlightbackground=self.cor2, highlightthickness=0.5, highlightcolor=self.cor2)
        self.frame_tela0.place(relx=0, rely=0, relwidth=1, relheight=1)
        # Abas Anotações
        self.abas_a = ttk.Notebook(self.aba4)
        self.abas_a.place(relx=0, rely=0, relwidth=1, relheight=1)
        # Tela 1
        self.aba_a_a = Frame(self.abas_a)
        self.aba_a_a.configure(background=self.cor1, bd=10)
        self.abas_a.add(self.aba_a_a, text='              Anotações sobre Alunos              ')
        # Tela 2
        self.aba_a_l = Frame(self.abas_a)
        self.aba_a_l.configure(background=self.cor1, bd=10)
        self.abas_a.add(self.aba_a_l, text='              Anotações sobre Livros              ')
    def anotacoes(self):

        ############# Parte de Criação da Função de Salvar #############

        def salva_alunos():
            if nome_e_a.get() == '' or t_anot_a.get("1.0",'end-1c') == '':
                messagebox.showinfo(title='ERRO!',message='Você não preencheu todas as informações')
            # Seleciona o Arquivo
            relatorio_aluno = Document('ARQUIVOS\ANOT\ALUNOS\BASE\BASE_anotacoes_.docx')
            # Edições dentro do Arquivo
            for paragrafo in relatorio_aluno.paragraphs:
                paragrafo.text = paragrafo.text.replace('X', nome_e_a.get())
                paragrafo.text = paragrafo.text.replace('W', self.format_data_retirada)
                paragrafo.text = paragrafo.text.replace('Y', t_anot_a.get("1.0",'end-1c'))
            # Salvar o Arquivo DOCX
            relatorio_aluno.save(f'ARQUIVOS\ANOT\ALUNOS\PRONTO\F_anotacoes_{(nome_e_a.get())}.docx')
            # Converte DOCX para PDF
            convert(f'ARQUIVOS\ANOT\ALUNOS\PRONTO\F_anotacoes_{(nome_e_a.get())}.docx',
                    f'ARQUIVOS\ANOT\ALUNOS\PRONTO\PDF\\anotacoes_{(nome_e_a.get())}.pdf')
            webbrowser.open(f'ARQUIVOS\ANOT\ALUNOS\PRONTO\PDF\\anotacoes_{(nome_e_a.get())}.pdf')
            nome_e_a.delete(0, END)
            t_anot_a.delete("1.0",'end-1c')
        def salva_livros():
            if nome_e_l.get() == '' or t_anot_l.get("1.0",'end-1c') == '':
                messagebox.showinfo(title='ERRO!',message='Você não preencheu todas as informações')
            # Seleciona o Arquivo
            relatorio_aluno = Document('ARQUIVOS\ANOT\LIVROS\BASE\BASE_anotacoes_.docx')                # Edições dentro do Arquivo
            for paragrafo in relatorio_aluno.paragraphs:
                paragrafo.text = paragrafo.text.replace('X', nome_e_l.get())
                paragrafo.text = paragrafo.text.replace('W', self.format_data_retirada)
                paragrafo.text = paragrafo.text.replace('Y', t_anot_l.get("1.0",'end-1c'))
            # Salvar o Arquivo em DOCX
            relatorio_aluno.save(f'ARQUIVOS\ANOT\LIVROS\PRONTO\F_anotacoes_{(nome_e_l.get())}.docx')
            # Converter DOCX para PDF
            convert(f'ARQUIVOS\ANOT\LIVROS\PRONTO\F_anotacoes_{(nome_e_l.get())}.docx',
                    f'ARQUIVOS\ANOT\LIVROS\PRONTO\PDF\\anotacoes_{(nome_e_l.get())}.pdf')
            webbrowser.open(f'ARQUIVOS\ANOT\LIVROS\PRONTO\PDF\\anotacoes_{(nome_e_l.get())}.pdf')
            nome_e_l.delete(0, END)
            t_anot_l.delete("1.0",'end-1c')

        ############# Parte de Anotações sobre alunos #############

        # Label de Titulo 1
        tit_a = Label(self.aba_a_a,text='Anotações de Alunos',bg=self.cor1,fg=self.cor11,
                                font=('Helvetica',15))
        tit_a.place(relx=0,rely=0,relheight=0.09,relwidth=1)
        # Label de Entrada de Nome 1
        nome_a = Label(self.aba_a_a,text='Nome do(a) Aluno(a):',bg=self.cor1,fg=self.cor11,
                                font=('Helvetica',10), anchor='w')
        nome_a.place(relx=0,rely=0.12,relheight=0.07,relwidth=1)
        # Entry de Entrada de Nome 1
        nome_e_a = Entry(self.aba_a_a,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor9,fg=self.cor7, insertontime='0')
        nome_e_a.place(relx=0.37,rely=0.11,relheight=0.09,relwidth=0.62)
        # Label de Entrada de Nome 1
        l_anot_a = Label(self.aba_a_a,text='Anotações',bg=self.cor1,fg=self.cor11,
                                font=('Helvetica',12, 'bold'), anchor='n')
        l_anot_a.place(relx=0,rely=0.24,relheight=0.08,relwidth=1)
        # Entrada de Anotações
        t_anot_a = Text(self.aba_a_a,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor9,fg=self.cor7)
        t_anot_a.place(relx=0.01,rely=0.34,relheight=0.65,relwidth=0.98)
        # Button de salvar a Anotação
        bt_salvar_a = Button(self.aba_a_a, command=salva_alunos,
                                    bg=self.cor6, bd=0, activebackground=self.cor6,
                                    highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                    text='Salvar', font=('Verdana', 10), cursor='hand2')
        bt_salvar_a.place(relx=0.79,rely=0.25,relheight=0.08,relwidth=0.2)

        ############# Parte de Anotações sobre alunos #############

        # Label de Titulo 2
        tit_l = Label(self.aba_a_l,text='Anotações de Livros',bg=self.cor1,fg=self.cor11,
                                font=('Helvetica',15))
        tit_l.place(relx=0,rely=0,relheight=0.09,relwidth=1)
        # Label de Entrada de Nome 2
        nome_l = Label(self.aba_a_l,text='Nome do Livro:',bg=self.cor1,fg=self.cor11,
                                font=('Helvetica',10), anchor='w')
        nome_l.place(relx=0,rely=0.12,relheight=0.07,relwidth=1)
        # Entry de Entrada de Nome 2
        nome_e_l = Entry(self.aba_a_l,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor9,fg=self.cor7, insertontime='0')
        nome_e_l.place(relx=0.27,rely=0.11,relheight=0.09,relwidth=0.72)
        # Label de Entrada de anotação 2
        l_anot_l = Label(self.aba_a_l,text='Anotações',bg=self.cor1,fg=self.cor11,
                                font=('Helvetica',12, 'bold'), anchor='n')
        l_anot_l.place(relx=0,rely=0.24,relheight=0.08,relwidth=1)
        # Entrada de Anotações
        t_anot_l = Text(self.aba_a_l,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor9,fg=self.cor7)
        t_anot_l.place(relx=0.01,rely=0.34,relheight=0.65,relwidth=0.98)
        # Button de salvar a Anotação
        bt_salvar_l = Button(self.aba_a_l, command=salva_livros,
                                    bg=self.cor6, bd=0, activebackground=self.cor6,
                                    highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                    text='Salvar', font=('Verdana', 10), cursor='hand2')
        bt_salvar_l.place(relx=0.79,rely=0.25,relheight=0.08,relwidth=0.2)
    def atrasados(self):
        # Criação da Treeview
        self.atrasa_bd = ttk.Treeview(self.frame_tela5, columns=('col1', 'col2', 'col3', 'col4', 'col5', 'col6'))
        self.atrasa_bd.place(relx=0, rely=0, relheight=1, relwidth=1)
        # Cabeçalho da Treeview
        self.atrasa_bd.heading('#0', text='')
        self.atrasa_bd.heading('#1', text='#')
        self.atrasa_bd.heading('#2', text='Aluno')
        self.atrasa_bd.heading('#3', text='')
        self.atrasa_bd.heading('#4', text='Livro')
        self.atrasa_bd.heading('#5', text='')
        self.atrasa_bd.heading('#6', text='Entrega')
        # Tamanho das colunas
        self.atrasa_bd.column('#0', width=0, stretch=0)
        self.atrasa_bd.column('#1', width=1, stretch=1)
        self.atrasa_bd.column('#2', width=110, stretch=1)
        self.atrasa_bd.column('#3', width=0, stretch=0)
        self.atrasa_bd.column('#4', width=130, stretch=1)
        self.atrasa_bd.column('#5', width=0, stretch=0)
        self.atrasa_bd.column('#6', width=80, stretch=1)
    def cadastro(self):
        # Frames
        # Frame de cima
        self.frame_cad1 = Frame(self.frame_tela3, bd=4, bg=self.cor1,
                                highlightbackground=self.cor2, highlightthickness=0.5, highlightcolor=self.cor2)
        self.frame_cad1.place(relx=0, rely=0, relwidth=1, relheight=0.32)
        # Frame do meio
        self.frame_cad2 = Frame(self.frame_tela3, bd=4, bg=self.cor1,
                                    highlightbackground=self.cor2, highlightthickness=0.5, highlightcolor=self.cor2)
        self.frame_cad2.place(relx=0, rely=0.34, relwidth=1, relheight=0.32)
        # Frame de baixo Esquerda
        self.frame_cad3 = Frame(self.frame_tela3, bd=4, bg=self.cor1,
                                    highlightbackground=self.cor2, highlightthickness=0.5, highlightcolor=self.cor2)
        self.frame_cad3.place(relx=0, rely=0.68, relwidth=0.495, relheight=0.32)
        # Frame de baixo Direita
        self.frame_cad4 = Frame(self.frame_tela3, bd=4, bg=self.cor1,
                                    highlightbackground=self.cor2, highlightthickness=0.5, highlightcolor=self.cor2)
        self.frame_cad4.place(relx=0.505, rely=0.68, relwidth=0.495, relheight=0.32)
        # Label
        # Label Aluno
        self.cad_label_aluno = Label(self.frame_cad1, text='Cadastro de alunos', bg=self.cor1, fg=self.cor3,
                                font=('Verdana', 10, 'bold'), anchor='w')
        self.cad_label_aluno.place(relx=0.02, rely=0.001, relheight=0.18, relwidth=0.35)
        # Label Livro
        self.cad_label_livro = Label(self.frame_cad2, text='Cadastro de livros', bg=self.cor1, fg=self.cor3,
                                font=('Verdana', 10, 'bold'), anchor='w')
        self.cad_label_livro.place(relx=0.02, rely=0.001, relheight=0.18, relwidth=0.35)
        # Label cadastro de escolas
        self.cad_label_escola = Label(self.frame_cad3, text='Cadastro de Escolas', bg=self.cor1, fg=self.cor3,
                                font=('Verdana', 12, 'bold'), anchor='n')
        self.cad_label_escola.place(relx=0, rely=0.1, relheight=0.3, relwidth=1)
        # Label cadastro de usuarios
        self.cad_label_usuarios = Label(self.frame_cad4, text='Cadastro de Usuários', bg=self.cor1, fg=self.cor3,
                                font=('Verdana', 12, 'bold'), anchor='n')
        self.cad_label_usuarios.place(relx=0, rely=0.1, relheight=0.3, relwidth=1)
        # Button
        # Button cadastrar Aluno
        self.cad_bt_aluno = Button(self.frame_cad1,
                                    bg=self.cor9, bd=0, activebackground=self.cor11, command=self.cad_aluno,
                                    highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                    text='Adicionar', font=('Verdana', 12), cursor='hand2')
        self.cad_bt_aluno.place(relx=0.75, rely=0.7, relheight=0.3, relwidth=0.25)
        # Button cadastrar Livro
        self.cad_bt_livro = Button(self.frame_cad2,
                                    bg=self.cor9, bd=0, activebackground=self.cor11, command=self.cad_livro,
                                    highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                    text='Adicionar', font=('Verdana', 12), cursor='hand2')
        self.cad_bt_livro.place(relx=0.75, rely=0.7, relheight=0.3, relwidth=0.25)
        # Button cadastrar escola
        self.cad_bt_escola = Button(self.frame_cad3, command=self.cad_escolas,
                                    bg=self.cor9, bd=0, activebackground=self.cor11,
                                    highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                    text='Adicionar', font=('Verdana', 12), cursor='hand2')
        self.cad_bt_escola.place(relx=0.25, rely=0.6, relheight=0.3, relwidth=0.5)
        # Button cadastrar usuario
        self.cad_bt_usuario = Button(self.frame_cad4, command=self.cad_usuarios,
                                    bg=self.cor9, bd=0, activebackground=self.cor11,
                                    highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                    text='Adicionar', font=('Verdana', 12), cursor='hand2')
        self.cad_bt_usuario.place(relx=0.25, rely=0.6, relheight=0.3, relwidth=0.5)
        # Text
        # Text Aluno
        label_aluno1 = Label(self.frame_cad1, text='Nome do Aluno, Ano escolar, Telefone,', bg=self.cor1, fg=self.cor5,
                                font=('Verdana', 12))
        label_aluno2 = Label(self.frame_cad1, text='Identificação, E-mail, Mãe, Pai', bg=self.cor1, fg=self.cor5,
                                font=('Verdana', 12))
        label_aluno1.place(relx=0.1, rely=0.3)
        label_aluno2.place(relx=0.08, rely=0.6)
        # Text Livro
        label_livro1 = Label(self.frame_cad2, text='Nome do Livro, Ano de Lançamento,', bg=self.cor1, fg=self.cor5,
                                font=('Verdana', 12))
        label_livro2 = Label(self.frame_cad2, text='Nome do Autor, Genero do Livro', bg=self.cor1, fg=self.cor5,
                                font=('Verdana', 12))
        label_livro1.place(relx=0.09, rely=0.3)
        label_livro2.place(relx=0.06, rely=0.6)
        
        # Text sem atribuição
    def botoes(self):
        
        ################### Tela 1 Inicio ###################

        self.print = PhotoImage(file='ARQUIVOS\IMAGENS\print.png') # Imagem de impressora
        # Botão Cadastrar aluno
        self.f1_bt_cadastrar = Button(self.frame_tela1,
                                    bg=self.cor9, bd=0, activebackground=self.cor11, command=self.cadastrar_f1,
                                    highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                    text='Cadastrar', font=('Verdana', 10), cursor='hand2')
        self.f1_bt_cadastrar.place(relx=0.555, rely=0.33, relheight=0.08, relwidth=0.2)
        # Botão imprimir cadastro
        self.f1_bt_imprimir1 = Button(self.frame_tela1,
                                    bg=self.cor9, bd=0, activebackground=self.cor9, command=self.geracomprovante,
                                    highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                    text='Imprimir', font=('Verdana', 10), cursor='hand2')
        self.f1_bt_imprimir1.place(relx=0.511, rely=0.92, relheight=0.08, relwidth=0.2)
        self.f1_bt_imprimir = Button(self.frame_tela1,
                                    bg=self.cor9, bd=0, activebackground=self.cor9, command=self.geracomprovante,
                                    highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                    image=self.print, font=('Verdana', 10), cursor='hand2')
        self.f1_bt_imprimir.place(relx=0.451, rely=0.92, relheight=0.08, relwidth=0.08)
        # Botão Limpar informações inseridas
        self.f1_bt_limpar = Button(self.frame_tela1, command=self.limpar_f1,
                                bg=self.cor8, bd=0, activebackground=self.cor11,
                                highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                text='Limpar', font=('Verdana', 10), cursor='hand2')
        self.f1_bt_limpar.place(relx=0.775, rely=0.33, relheight=0.08, relwidth=0.2)
        # Botão excluir cadastro
        self.f1_bt_excluir = Button(self.frame_tela1, command=self.excluir_f1,
                                bg=self.cor8, bd=0, activebackground=self.cor11,
                                highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                text='Excluir', font=('Verdana', 10), cursor='hand2')
        self.f1_bt_excluir.place(relx=0.02, rely=0.92, relheight=0.08, relwidth=0.2)
        # Botão alterar cadastro
        self.f1_bt_alterar = Button(self.frame_tela1, command=self.alterar_f1,
                                bg=self.cor6, bd=0, activebackground=self.cor11,
                                highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                text='Alterar', font=('Verdana', 10), cursor='hand2')
        self.f1_bt_alterar.place(relx=0.235, rely=0.92, relheight=0.08, relwidth=0.2)
        
        ################### Tela 2 Consulta ###################
    def entry(self):

        ################### Tela 1 Inicio ###################

        ##### Entrys das datas (não aparecem na tela principal)
        # Criação e inserção da data de retirada
        self.entry_data_retirada = Entry(self.frame_tela1)
        self.entry_data_retirada.insert(0, self.format_data_retirada)
        # Criação e inserção da data de entrega
        self.entry_data_entrega = Entry(self.frame_tela1)
        self.entry_data_entrega.insert(0, self.format_data_entrega)
        ######################################################
        # Entry de codigo tela 1
        self.f1_entry_cod = Entry(self.frame_tela1,
                                    bg=self.cor1, bd=3, highlightbackground=self.cor1, relief=FLAT, justify='center',
                                    highlightthickness=2, highlightcolor=self.cor1,fg=self.cor3, insertontime='0', cursor='arrow')
        self.f1_entry_cod.place(relx=0.45, rely=0.07, relheight=0.08, relwidth=0.17)
        # Entry de nome de aluno tela 1
        self.f1_entry_aluno = Entry(self.frame_tela1,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10,fg=self.cor2, insertontime='0')
        self.f1_entry_aluno.place(relx=0.02, rely=0.07, relheight=0.08, relwidth=0.4)
        # Entry de nome de livro tela 1
        self.f1_entry_livro = Entry(self.frame_tela1,
                                    bg=self.cor11, bd=0, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor10, fg=self.cor2, insertontime='0')
        self.f1_entry_livro.place(relx=0.02, rely=0.25, relheight=0.08, relwidth=0.4)

        ##################################################################################
        #######################Drop-Down e Entry Drop-Down################################
        ##################################################################################
        # Variavel do Drop-Down de serie do aluno tela 1
        self.f1_anos_escolares = ('1° Ano', '2° Ano', '3° Ano', '4° Ano', '5° Ano', '6° Ano', '7° Ano', '8° Ano', '9° Ano')
        self.f1_variavel_serie = StringVar()
        self.f1_variavel_serie.set('Selecionar')
        # Criação do Drop-Down de serie do aluno tela 1
        self.f1_drop_serie = OptionMenu(self.frame_tela1, self.f1_variavel_serie, *self.f1_anos_escolares,
                                    command=self.optionmenu)
        self.f1_drop_serie.place(relx=0.65, rely=0.07, relheight=0.08, relwidth=0.29)
        self.f1_drop_serie.config(bg=self.cor11, bd=0, highlightbackground=self.cor10,  relief=FLAT,
                                fg=self.cor7, font=('Verdana', 11), indicatoron='0', direction='below',
                                cursor='hand2')
        # Entry Serie do aluno
        self.f1_entry_serie = Entry(self.frame_tela1,
                                    bg=self.cor1, bd=0, highlightbackground=self.cor1, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor1, fg='green', insertontime='0', cursor='arrow')
        self.f1_entry_serie.place(relx=0.6, rely=0.15, relheight=0.08, relwidth=0.390)

        ################### Tela 2 Consulta ###################

        # Entry código da Tela de informações
        self.e_c = Entry(self.root)
        # Entry nome da Tela de informações
        self.e_n = Entry(self.root)
        # Entry serie da Tela de informações
        self.e_s = Entry(self.root)
        # Entry livro da Tela de informações
        self.e_l = Entry(self.root)
        # Entry data1 da Tela de informações
        self.e_d1 = Entry(self.root)
        # Entry data2 da Tela de informações
        self.e_d2 = Entry(self.root)
        # Entry pesq aluno
        self.f2_entry_aluno = Entry(self.frame_tela2,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor9,fg=self.cor7, insertontime='0')
        self.f2_entry_aluno.place(relx=0.01, rely=0.08, relheight=0.08, relwidth=0.24)
        # Entry pesq serie
        self.f2_entry_serie = Entry(self.frame_tela2,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor9,fg=self.cor7, insertontime='0')
        self.f2_entry_serie.place(relx=0.27, rely=0.08, relheight=0.08, relwidth=0.24)
        # Entry pesq livro
        self.f2_entry_livro = Entry(self.frame_tela2,
                                    bg=self.cor11, bd=3, highlightbackground=self.cor6, relief=FLAT,
                                    highlightthickness=2, highlightcolor=self.cor9, fg=self.cor7, insertontime='0')
        self.f2_entry_livro.place(relx=0.528, rely=0.08, relheight=0.08, relwidth=0.24)
        # Button pesquisa
        self.f2_bt_pesquisa = Button(self.frame_tela2, command=self.pesquisar_f2,
                                    bg=self.cor6, bd=0, activebackground=self.cor9, activeforeground=self.cor3,
                                    highlightbackground=self.cor2, highlightthickness=1, fg=self.cor3,
                                    text='Pesquisar', font=('Verdana', 10), cursor='hand2')
        self.f2_bt_pesquisa.place(relx=0.79, rely=0.08, relheight=0.08, relwidth=0.19)

        ################### Tela 3 Cadastro ###################
    def treeview_f1(self):
        # Criação da Treeview
        self.inser_bd = ttk.Treeview(self.frame_tela1, columns=('col1', 'col2', 'col3', 'col4', 'col5', 'col6'))
        self.inser_bd.place(relx=0.01, rely=0.42, relheight=0.49, relwidth=0.91)
        # Cabeçalho da Treeview
        self.inser_bd.heading('#0', text='')
        self.inser_bd.heading('#1', text='#')
        self.inser_bd.heading('#2', text='Aluno')
        self.inser_bd.heading('#3', text='Ano')
        self.inser_bd.heading('#4', text='Livro')
        self.inser_bd.heading('#5', text='Retirada')
        self.inser_bd.heading('#6', text='Entrega')
        # Tamanho das colunas
        self.inser_bd.column('#0', width=0, stretch=0)
        self.inser_bd.column('#1', width=1, stretch=1)
        self.inser_bd.column('#2', width=80, stretch=1)
        self.inser_bd.column('#3', width=50, stretch=1)
        self.inser_bd.column('#4', width=100, stretch=1)
        self.inser_bd.column('#5', width=0, stretch=0)
        self.inser_bd.column('#6', width=50, stretch=1)
        # Personalizando as cores da Treeview
        treeview_style = ttk.Style()
        treeview_style.theme_use('default') # clam , default , alt , vista
        treeview_style.configure('Treeview',
                                    background=self.cor1,
                                    foreground=self.cor3,
                                    rowheight=25,
                                    fieldbackground=self.cor4)
        treeview_style.map('Treeview',
                                background=[('selected', self.cor9)])                               
        # Barra de rolagem Vertical
        self.scroll_inser_V = Scrollbar(self.frame_tela1, orient='vertical')
        self.inser_bd.configure(yscroll=self.scroll_inser_V.set)
        self.scroll_inser_V.place(relx=0.92, rely=0.423, relheight=0.488, relwidth=0.06)
        self.scroll_inser_V.config(background=self.cor1, troughcolor=self.cor1, activebackground=self.cor2)
        # Chamar a função DB Click
        self.inser_bd.bind('<Double-1>', self.db_click)
    def treeview_f2(self):
        # Criação da Treeview
        self.pesq_bd = ttk.Treeview(self.frame_tela2, columns=('col1', 'col2', 'col3', 'col4', 'col5', 'col6'))
        self.pesq_bd.place(relx=0.01, rely=0.2, relheight=0.78, relwidth=0.95)
        # Cabeçalho da Treeview
        self.pesq_bd.heading('#0', text='')
        self.pesq_bd.heading('#1', text='#')
        self.pesq_bd.heading('#2', text='Aluno')
        self.pesq_bd.heading('#3', text='Ano')
        self.pesq_bd.heading('#4', text='Livro')
        self.pesq_bd.heading('#5', text='Retirada')
        self.pesq_bd.heading('#6', text='Entrega')
        # Tamanho das colunas
        self.pesq_bd.column('#0', width=0, stretch=0)
        self.pesq_bd.column('#1', width=0, stretch=0)
        self.pesq_bd.column('#2', width=150, stretch=1)
        self.pesq_bd.column('#3', width=80, stretch=1)
        self.pesq_bd.column('#4', width=150, stretch=1)
        self.pesq_bd.column('#5', width=0, stretch=0)
        self.pesq_bd.column('#6', width=0, stretch=0)                         
        # Barra de rolagem Vertical
        self.scroll_pesq_V = Scrollbar(self.frame_tela2, orient='vertical')
        self.pesq_bd.configure(yscroll=self.scroll_pesq_V.set)
        self.scroll_pesq_V.place(relx=0.96, rely=0.202, relheight=0.7777, relwidth=0.03)
        self.scroll_pesq_V.config(background=self.cor1, troughcolor=self.cor1, activebackground=self.cor2)
        self.pesq_bd.bind('<Double-1>', self.db_click_t2)
    def date_today(self):
        # Busca da Data atual
        self.data_hoje = date.today()
        # Soma de 7 dias na Data Atual
        self.data_t_delta = timedelta(7)
        self.data_para_entrega = self.data_hoje + self.data_t_delta
        # Formatação das Datas
        self.format_data_retirada = self.data_hoje.strftime("%d/%m/%Y")
        self.format_data_entrega = self.data_para_entrega.strftime("%d/%m/%Y")
        # Mostrador da data Atual
        self.data_atual = Label(self.frame_tela1, text=self.format_data_retirada, bg=self.cor1, fg='white',
                                font=('Verdana', 10))
    def msg_box(self):
        ##############################################
        #########_________variaveis_________##########
        #######_________Tela de inicio_________#######
        msg_bt_limpar = 'Limpar informações'
        msg_bt_alterar = 'Alterar informações'
        msg_bt_excluir = 'Excluir informações'
        msg_bt_cadastrar = 'Cadastrar informações'
        msg_bt_imprimir = 'Imprimir comprovante'
        #######________Tela de Consulta_________#######
        msg_et_nome = 'Insira o nome do aluno'
        msg_et_livro = 'Insira o nome do livro'
        msg_et_serie = 'Insira o ano do aluno'
        msg_bt_pesquisar = 'Pesquisar'
        ########_________Tela de Login_________########
        msg_l_login = 'Tela de Login'
        msg_nome_login = 'Insira seu nome de usuário'
        msg_senha_login = 'Insira sua senha'
        msg_m_senha_login = 'Selecione para tornar sua senha visivel'
        msg_b_entrar = 'Entrar com suas credenciais'
        msg_b_cancel = 'Cancelar entrada e sair do aplicativo'
        ########_________Tela de Cadastro_________########
        msg_add_aluno = 'Clique para adicionar um aluno'
        msg_add_livro = 'Clique para adicionar um Livro'
        msg_add_escola = 'Clique para adicionar uma Escola'
        msg_add_usuario = 'Clique para adicionar um Usuário'
        ###############################################
        ########________Tela de Inicio_________########
        ###############################################
        # msg BT Limpar Frame 1
        self.msg_limpar = Pmw.Balloon(self.frame_tela1)
        self.msg_limpar.bind(self.f1_bt_limpar, msg_bt_limpar)
        # msg BT Cadastrar Frame 1
        self.msg_cadastrar = Pmw.Balloon(self.frame_tela1)
        self.msg_cadastrar.bind(self.f1_bt_cadastrar, msg_bt_cadastrar)
        # msg BT Alterar Frame 1
        self.msg_alterar = Pmw.Balloon(self.frame_tela1)
        self.msg_alterar.bind(self.f1_bt_alterar, msg_bt_alterar)
        # msg BT Excluir Frame 1
        self.msg_excluir = Pmw.Balloon(self.frame_tela1)
        self.msg_excluir.bind(self.f1_bt_excluir, msg_bt_excluir)
        # msg BT Imprimir Frame 1
        self.msg_imprimir = Pmw.Balloon(self.frame_tela1)
        self.msg_imprimir1 = Pmw.Balloon(self.frame_tela1)
        self.msg_imprimir.bind(self.f1_bt_imprimir, msg_bt_imprimir)
        self.msg_imprimir1.bind(self.f1_bt_imprimir1, msg_bt_imprimir)
        # msg Entry Nome Frame 1
        self.msg_nome1 = Pmw.Balloon(self.frame_tela1)
        self.msg_nome1.bind(self.f1_entry_aluno, msg_et_nome)
        # msg Entry Livro Frame 1
        self.msg_livro1 = Pmw.Balloon(self.frame_tela1)
        self.msg_livro1.bind(self.f1_entry_livro, msg_et_livro)
        # msg Menu serie Frame 1
        self.msg_serie1 = Pmw.Balloon(self.frame_tela1)
        self.msg_serie1.bind(self.f1_drop_serie, msg_et_serie)
        #######________Tela de Consulta_________#######
        # msg BT pesquisar Frame 2
        self.msg_pesq = Pmw.Balloon(self.frame_tela2)
        self.msg_pesq.bind(self.f2_bt_pesquisa, msg_bt_pesquisar)
        # msg Entry Nome Frame 2
        self.msg_nome2 = Pmw.Balloon(self.frame_tela2)
        self.msg_nome2.bind(self.f2_entry_aluno, msg_et_nome)
        # msg Entry Livro Frame 2
        self.msg_livro2 = Pmw.Balloon(self.frame_tela2)
        self.msg_livro2.bind(self.f2_entry_livro, msg_et_livro)
        # msg Entry serie Frame 2
        self.msg_serie2 = Pmw.Balloon(self.frame_tela2)
        self.msg_serie2.bind(self.f2_entry_serie, msg_et_serie)
        #######________Tela de Login_________#######
        # msg Label Login Frame login
        self.msg_l_login = Pmw.Balloon(self.frame_tela0)
        self.msg_l_login.bind(self.l_label_tit, msg_l_login)
        # msg Entry usuario Frame login
        self.msg_nome_login = Pmw.Balloon(self.frame_tela0)
        self.msg_nome_login.bind(self.l_entry_usu, msg_nome_login)
        # msg Entry senha Frame login
        self.msg_senha_login = Pmw.Balloon(self.frame_tela0)
        self.msg_senha_login.bind(self.l_entry_sen1, msg_senha_login)
        # msg de mostrar senha Frame login
        self.msg_m_senha_login = Pmw.Balloon(self.frame_tela0)
        self.msg_m_senha_login.bind(self.l_ck_senha, msg_m_senha_login)
        self.msg_m_senha_login.bind(self.l_label_ck, msg_m_senha_login)
        # msg Button Entrar Frame login
        self.msg_b_entrar = Pmw.Balloon(self.frame_tela0)
        self.msg_b_entrar.bind(self.l_bt_login, msg_b_entrar)
        # msg Button Cancelar Frame login
        self.msg_b_cancel = Pmw.Balloon(self.frame_tela0)
        self.msg_b_cancel.bind(self.l_bt_can, msg_b_cancel)
        # msg Button Adicionar Aluno Frame Cadastro
        self.msg_add_aluno = Pmw.Balloon(self.frame_cad1)
        self.msg_add_aluno.bind(self.cad_bt_aluno, msg_add_aluno)
        # msg Button Adicionar Livro Frame Cadastro
        self.msg_add_livro = Pmw.Balloon(self.frame_cad2)
        self.msg_add_livro.bind(self.cad_bt_livro, msg_add_livro)
        # msg Button Adicionar Escola Frame Cadastro
        self.msg_add_livro = Pmw.Balloon(self.frame_cad3)
        self.msg_add_livro.bind(self.cad_bt_escola, msg_add_escola)
        # msg Button Adicionar Usuário Frame Cadastro
        self.msg_add_livro = Pmw.Balloon(self.frame_cad4)
        self.msg_add_livro.bind(self.cad_bt_usuario, msg_add_usuario)

Application()